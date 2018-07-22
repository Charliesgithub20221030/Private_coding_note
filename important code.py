

#### --This the key code in super data science


from sklearn.linear_model import LinearRegression
regressor=LinearRegression()
regressor.fit(x,y)


from sklearn.preprocessing import PolynomialFeatures
poly=PolynomialFeatures(degree=4)
x_poly=poly.fit_transform(x)

from sklearn.svm import SVR
regressor=SVR()
regressor.fit(x,y)

from sklearn.tree import DecisionTreeRegressor
regressor=DecisionTreeRegressor()
regressor.fit(x,y)

from sklearn.ensemble import RandomForeseRegressor
regressor=RandomForeseRegressor()
regressor.fit(x,y)

from sklearn.linear_model import LogisticRegression
classifier=LogsticRegression()
classifier.fit(x,y)

from sklearn.neighbors import KNeighborsClassifier
classifier=KNeighborsClassifier()
classifier.fit(x,y)


from sklearn.svm import SVC
classifier= SVC
classifier.fit(x,y)


###  #related code

##--1
enumerate(list)
a=['a','b','c','d']
#普通的for 
for i in a:
    print(i)
#a
#b
#c
#d
#例舉的for
for n,i in enumerate(a):
    print(n,i)
#0 a
#1 b
#2 c
#3 d

###--2
#字串中置入字元(字串前加上f)
waitTime = 2018/2/2
print(f"At time:{clawerTime} , waitTime:{waitTime}")
#另外的方法
i=0
e=9
'%d and %s' % (i, e)


###--3
#讀取Docx
From docx import Document
F=Document(‘filename’)
For i In f.paragraphs:
	Print(i.text)
#完整版
import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText) #用join前面的字元連接所有list成員

print(getText('台灣行銷研究.docx'))



###結束python
import sys
sys.exit(1) 
import os
os._exit(1)

###groupby
#直接看網路
https://ithelp.ithome.com.tw/articles/10194027

###std
#numpy.std() 求标准差的时候默认是除以 n 的，即是有偏的，np.std无偏样本标准差方式为 ddof = 1； 
#pandas.std() 默认是除以n-1 的，即是无偏的，如果想和numpy.std() 一样有偏，需要加上参数ddof=0 ，即pandas.std(ddof=0) ；DataFrame的describe()中就包含有std()；
#
#>>> a
#array([0, 1, 2, 3, 4, 5, 6, 7, 8, 9])
#>>> np.std(a, ddof = 1)
#3.0276503540974917
#>>> np.sqrt(((a - np.mean(a)) ** 2).sum() / (a.size - 1))
#3.0276503540974917
#>>> np.sqrt(( a.var() * a.size) / (a.size - 1))
#3.0276503540974917


#函數定義時就已經宣告a=[]，所以a不會被清空
def b(x,a=[]):
    
    a.append(x)
    return a
    
print(b('asd'))
print(b('qwe'))
print(b('zxc'))
#必須寫成這樣
def b(x):
    a=[]
    a.append(x)
    return a


#引數用**，產生字典
def d(**arg):
    return (arg)
a=d(k1='kk',k2='kkk',k3='kkkk')
print(a)

##文件撰寫
def print_more(*a):
    'print more than one argument'
    print(a)
#查看資訊
help(print_more)
#單純查看資訊
print(print_more.__doc__)#注意，共有四個 '_'


##closure:由函數生成的函數，其傳回值為一個函數，
#         所以變數可以被當成這個函數的複製品執行
def a1(say):
    def a2():
        return 'I am a2 , this is saying : %s'%say
    return a2
a=a1('Hello Charlie')
print(a())

#lambda 的使用方式
voice=['wo','meu','mo','bi','gang']
def v(l,func):
    for i in l:
        print(func(i))
v(voice,lambda word: word.capitalize() + '!')



#yield 表示生成器，必須用迭代方式取出
def y(n):
    a=1
    while a<=n:
        yield a
        a+=1
p=y(20)
for i in p:
    print(i)


#裝飾器的使用
def document_it(func):
    def new_func(*arg,**kwargs):
        result=func(*arg,**kwargs)
        print('Arguments : ',arg)
        print('Dictionary : ',kwargs)
        print('Result : ', result)
        return result
    return new_func

def add_it(*arg):
    return sum(arg)
doc=document_it(add_it) #第一種呼叫方法，放在變數
print(doc(1,2,3,4,5,6,7,8,9))

@document_it            #第二種呼叫方法，加上@，而且可以多個，但是順序有差
def add_it(*arg):       #比較下面的(比較接近)，會先執行
    return sum(arg)

print(add_it(1,2,3,4,5,6,7,8,9))

def square_it(func):
    def new_func(*arg,**kwargs):
        result=func(*arg,**kwargs)
        print('Arguments : ',arg)
        print('Dictionary : ',kwargs)
        print('Result : ',result**2)
    return new_func

@square_it   #雙重decorater 範例
@document_it
def add_it(*arg):
    return sum(arg)
print(add_it(1,2,3,4,5,6,7,8,9))


#numpy 的生成工具
np.random.randn(1000)
np.arange(1000)


#enumerate 的用途；迭代，並傳回編號(從1開始)
li=[i for i in range(50,71)]
for i,j in enumerate(li,1):
    print(i,j)
    
    
    
#default dict 只有有key，就可以預設value的字典
from collections import defaultdict
di=defaultdict(int) #後面必須附上轉換函數
ds=defaultdict(str)
ds['A']
ds['B']=0
dd=defaultdict(dict)


    
#印出字典鍵值

for i,j in zip('ABCDEFG',range(7)): #範例字典
    z[i]=j
    
for i,j in z.items():
    print(i,j)
    
    
#計算神器 collections -----Counter
from collections import Counter

s='''I've been reading books of old
The legends and the myths
Achilles and his gold
Hercules and his gifts
Spiderman's control
And Batman with his fists
And clearly I don't see myself upon that list
But she said, where'd you wanna go?
How much you wanna risk?
I'm not looking for somebody
With some superhuman gifts
Some superhero
Some fairytale bliss
Just something I can turn to
Somebody I can kiss
I want something just like this
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Oh, I want something just like this
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Oh, I want something just like this
I want something just like this
I've been reading books of old
The legends and the myths
The testaments they told
The moon and its eclipse
And Superman unrolls
A suit before he lifts
But I'm not the kind of person that it fits
She said, where'd you wanna go?
How much you wanna risk?
I'm not looking for somebody
With some superhuman gifts
Some superhero
Some fairytale bliss
Just something I can turn to
Somebody I can miss
I want something just like this
I want something just like this
I want something just like this
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Oh, I want something just like this
Doo-doo-doo, doo-doo
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Where'd you wanna go?
How much you wanna risk?
I'm not looking for somebody
With some superhuman gifts
Some superhero
Some fairytale bliss
Just something I can turn to
Somebody I can kiss
I want something just like this
Oh, I want something just like this
Oh, I want something just like this
Oh, I want something just like this'''
s=s.replace('\n','')
lists=s.split(' ')

c=Counter(lists)  #直接將list依照詞頻轉換字典
c.most_common(1)  #顯示前1高詞頻的鍵值對

#利用counter 來做集合運算
s2='''Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Oh, I want something just like this
Doo-doo-doo, doo-doo
Doo-doo-doo, doo-doo-doo
Doo-doo-doo, doo-doo-doo
Where'd you wanna go?
How much you wanna risk?The default factory is called without arguments to produce a new value when a key is not present, in __getitem__ only. A defaultdict compares equal to a dict with the same items. All remaining arguments are treated the same as if they were passed to the dict constructor, including keyword arguments.
'''
c2=Counter(s2)
c-c2
c&c2
c|c2

#一樣產生字典，但是利用OrderedDict 產生的字典有明確的順序，一般字典比較不確定
from collections import OrderedDict

# deque (讀作deck) 可以同時使用stack 和 queue
from collections import deque
c=deque('asdfghjkl')
c.pop()
c.popleft()


#python沒有字串 reverse(但是有list)，可是可以利用slice的方法反轉
'asdfghjkl'[::-1]

#### itertools 迭代神器

#chain 
from itertools import chain #無視容器邊界，直接迭代
a=range(10)
b=list('asdfghjk')
for i in chain(a,b):  #很適合用於將不同類的大型資料整合
    print(i)
    


#cycle 無限循環迭代
from itertools import cycle
n=0
a=[1,2,3]
for i in cycle(a):
    print(i , end = ' ')
    n+=1
    if n==100:
        break
    
    
#accumulate 累加迭代
from itertools import accumulate
a=range(1,21)
for i in accumulate(a):
    print(i)



####高級輸出器 pprint
#可以將輸出依照元素排列  ---- 其實我還看不出差別
from pprint import pprint
d=dict(['a1','b2','c3','d4','e5','f6','g7'])
print(d)
pprint(d)



####解決中文編碼問題---chardet-----查詢字串的encoding
import chardet


#可以用下列方法取得文字檔的編碼，結果會存在code裡面，就可以用來設定encoding
fileName='自己的檔案自己輸....txt'
with open(fileName,'rb') as f:
    fileLines=f.readlines()
for i in fileLines:
    getCode=chardet.detect(i)['encoding']
    if getCode!=None:
        code=getCode
        break
#然後利用decode(code)來解碼 binary ---可以 print出來試試看
for i in fileLines:
    print(i.decode(code))


#### real python trick---------------------------------------------
#讓dict依照values排序
sorted(d.items(),key=lambda x: x[1])


#如何合併兩個字典
d=dict()
for i,j in zip('ABCD',range(1,5)):
    d[i]=j
dd=dict()
for i,j in zip('DEFG',range(11,15)):
    dd[i]=j
z={**d,**dd}


#------------------------------------------------------------------


z={**d,**dd}
#Out[79]: {'A': 1, 'B': 2, 'C': 3, 'D': 11, 'E': 12, 'F': 13, 'G': 14}






#scatter範例
#因為scatter只能放X,Y兩個數，所以要用疊加的方式
data = pd.DataFrame(
    np.random.randn(1000,4),#多了一個引數，可以變成多維
    index=np.arange(1000),
    columns=list("ABCD")
    )
data.cumsum()
graph=data.plot.scatter('A','B',color='red',label='B')
data.plot.scatter('A','C',color='blue',label='C',ax=graph)  #注意要加ax=graph，
data.plot.scatter('A','D',color='green',label='D',ax=graph) #不然會分成四張圖
plt.show()



#多程序
import threading as td
def doSomething(a,b):
    print(a+b)
t1=td.Thread(target=doSomething,args=(1,2))
t1.start()
t1.join()
#多核
import multiprocessing as mp
def doSomething(a,b):
    print(a+b)
if __name__=='__main__':
    m1=mp.Process(target=doSomething,args=(1,2))
    m1.start()
    m1.join()




#re
import re 

a=re.search('a.c','abcdd')
a=re.match('a.c','aaabc')#match 是從頭開始配對，所以這樣找不到
#沒有東西就傳回 NoneType

#其實要做其他動作，要先把規則放在compile 裡面，下面是 substitute
reg=re.compile('[a-v]*')
reg.sub('','qwertyuioabasdfghjkab') #取代成空字串-清理資料很有用

re.escape('1234@gmail.com') #Escape all the characters in pattern except ASCII
                            # letters, numbers and '_'.


#處理missing data時要注意使用行資料，或列資料
#如果是 np.array ，要用reshape(-1,1)(行，長條的)，reshape(1,-1)(列，寬的)
#處理過後才能在處理missing data
#還有一點，取出dataframe的行或列時，要用.values，取出的才是array，否則是pandas Series
df =  pd.read_csv("Melbourne_housing_full.csv")
y=df['Price'].values
from sklearn.preprocessing import Imputer
ip=Imputer(missing_values='NaN',strategy='mean',axis=0)
y=ip.fit_transform(y.reshape(-1,1))

###---Pandas
#測試專用資料
df=pd.DataFrame({'AA':np.arange(30,61),'BB':np.arange(31)})
ddf=pd.DataFrame({'A':[1],'B':['abc']})



#重設 index(當增刪列的時候，如果還要用index的方式用迴圈搜尋，必須要重設)
df.reset_index(drop=True)
df=df.reset_index(drop=True)#只會傳回值，必須用接收的


#取得資料框中特定儲存格(cell)
df.loc[0][1]

#取得特定列
df.loc[0]#找格的邏輯就是，先找到列，然後把列當成陣列，再找到特定的cell



#dataframe 可以 append ，append 完还是一個 dataframe
#可是 append 完 不會傳回，所以要用原始DF接收
df=df.append(df2)


#drop row by index
df=df.drop([1,2])
#by name
#--

#drop colums by index
df=df.drop(df.columns[1])
#by name
df=df.drop(df['name1','name2'],axis=1)


#pandas dataframe 的神器 dataframe.dropna()


#在终端窗口中输出图片的命令是：
# %matplotlib inline

#在图片窗口表现图片的命令是：
# %matplotlib qt   -------必須先把 IPthon->graphics 設成 Qt5 ->重開

#不顯示科學記號的方法
import numpy as np
np.set_printoptions(suppress=True)




#預設是engine='c'，可是如果是中文可能會出現問題，所以要改變引數
data1=pd.read_csv('乙太幣.csv',engine='python')

#轉換dataframe型態
 dataframe.astype(int)

#row names of data frame
dataframe.iloc[[0,:]].index.tolist()
dataframe.axes[0].tolist()

#dataframe 合併
#轉列表，創造新DF
df=pd.DataFrame({'Date':list(Date),
                 'Score':list(total_score)})
#append
df1.append([df2,df3...])

#dataframe---concatenating 
# axis=0:列合併
# axis=1:欄合併
#ignore_index=True :表示，重新標註列名，否則會出現 0 1 2 0 1 2 3 0 1 2...的情形
#join引數中，'inner'表示[無此欄，刪除]，'outer'表示[無此欄仍加入，但是不足者，以NaN填充]
allData=pd.concat([dataframe1,dataframe2,dataframe3],axis=1,ignore_index=True,join=)
allData=pd.concat([dataframe1,dataframe2,dataframe3],axis=0,join='inner'/'outer')                        

#merge-內建的合併方法是inner(在這裡不叫 join ，叫做 how ='inner')
new_df=pd.merge(df1,df2,on='a column name'/['col1','col2',....])#橫向合併,on表示相同的欄



#pandas 內鍵繪圖功能範例
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt

data = pd.Series(np.random.randn(1000),index=np.arange(1000))#生成以 np.arange 為列名的亂數 dataframe
data.cumsum()#這是累加
data.plot()
plt.show()

#四欄繪圖
data = pd.DataFrame(
    np.random.randn(1000,4),#多了一個引數，可以變成多維
    index=np.arange(1000),
    columns=list("ABCD")
    )
data.cumsum()
data.plot()
plt.show()


#除了plot，我经常会用到还有scatter，这个会显示散点图，首先给大家说一下在 pandas 中有多少种方法
#
#bar
#hist
#box
#kde
#area
#scatter
#hexbin



%matplotlib qt
import pandas as pd
import numpy as np
import seaborn as sns
df =  pd.read_csv("Melbourne_housing_full.csv")
y=df['Price'].values
X=df['BuildingArea'].values
df['Price']=y
df['BuildingArea']=X
from sklearn.preprocessing import Imputer
ip=Imputer(missing_values='NaN',strategy='mean',axis=0)
y=ip.fit_transform(y.reshape(-1,1))
X=ip.fit_transform(X.reshape(-1,1))
sns.lmplot(x='Distance',y='Price',data=df,hue='Regionname',col='Regionname',col_wrap=4,order=2,size=3)
np.array(df['Price']).reshape(1,-1)
